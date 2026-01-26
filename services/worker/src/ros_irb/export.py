from __future__ import annotations

from pathlib import Path
from typing import Sequence

from ros_irb.generate_irb_request import IRBDraft
from ros_irb.irb_questions import IRBQuestion, IRB_QUESTIONS
from ros_irb.phi_guard import redact_phi


def export_docx(
    draft: IRBDraft,
    output_path: Path,
    *,
    questions: Sequence[IRBQuestion] = IRB_QUESTIONS,
    redact: bool = True,
) -> Path:
    """
    Export IRB draft to a Word document (DOCX format).

    Creates a structured DOCX document with headings and paragraphs from the
    IRB draft content, applying PHI redaction if requested.

    Parameters
    ----------
    draft : IRBDraft
        The IRB draft object containing study title, answers, and metadata.
    output_path : Path
        Path where the DOCX file will be saved. Parent directories are created
        automatically if they do not exist.
    questions : Sequence[IRBQuestion], optional
        List of IRB questions to include in the export. Defaults to IRB_QUESTIONS.
    redact : bool, optional
        If True (default), applies PHI redaction to all text content using
        the redact_phi() function. PHI patterns (emails, phone numbers, etc.)
        are replaced with "[REDACTED]".

    Returns
    -------
    Path
        The path to the created DOCX file.

    Notes
    -----
    - Requires python-docx package (optional dependency)
    - PHI redaction is applied by default for compliance with governance policies
    - Document structure: Title (H1), Generated date, Literature Context (H2 if present),
      then each question with its prompt and answer

    Examples
    --------
    >>> from ros_irb.export import export_docx
    >>> path = export_docx(draft, Path("output/draft.docx"), redact=True)
    """
    from docx import Document  # python-docx

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(f"IRB Draft: {draft.study_title}", level=1)
    doc.add_paragraph(f"Generated: {draft.created_at_iso}")

    if draft.literature_summary:
        doc.add_heading("Literature Context (Auto-generated)", level=2)
        doc.add_paragraph(redact_phi(draft.literature_summary) if redact else draft.literature_summary)

    for q in questions:
        doc.add_heading(q.title, level=2)
        doc.add_paragraph(q.prompt)
        ans = (draft.answers.get(q.category) or "").strip() or "[No response provided yet]"
        doc.add_paragraph(redact_phi(ans) if redact else ans)

    doc.save(str(output_path))
    return output_path


def export_pdf(
    draft: IRBDraft,
    output_path: Path,
    *,
    questions: Sequence[IRBQuestion] = IRB_QUESTIONS,
    redact: bool = True,
) -> Path:
    """
    Export IRB draft to a PDF document.

    Creates a PDF document with the IRB draft content using reportlab,
    applying PHI redaction if requested. Uses text wrapping to handle
    long lines properly without truncation.

    Parameters
    ----------
    draft : IRBDraft
        The IRB draft object containing study title, answers, and metadata.
    output_path : Path
        Path where the PDF file will be saved. Parent directories are created
        automatically if they do not exist.
    questions : Sequence[IRBQuestion], optional
        List of IRB questions to include in the export. Defaults to IRB_QUESTIONS.
    redact : bool, optional
        If True (default), applies PHI redaction to all text content using
        the redact_phi() function. PHI patterns (emails, phone numbers, etc.)
        are replaced with "[REDACTED]".

    Returns
    -------
    Path
        The path to the created PDF file.

    Notes
    -----
    - Requires reportlab package (optional dependency)
    - PHI redaction is applied by default for compliance with governance policies
    - Page breaks are inserted automatically when content exceeds page height
    - Long lines are wrapped at word boundaries (max 90 characters per line)
    - Document uses Letter page size (8.5" x 11")

    Examples
    --------
    >>> from ros_irb.export import export_pdf
    >>> path = export_pdf(draft, Path("output/draft.pdf"), redact=True)
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter

    y = height - 72
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, y, f"IRB Draft: {draft.study_title}")
    y -= 24
    c.setFont("Helvetica", 10)
    c.drawString(72, y, f"Generated: {draft.created_at_iso}")
    y -= 24

    def wrap_text(text: str, max_chars: int = 90) -> list[str]:
        """Wrap text at word boundaries to fit within max_chars per line."""
        import textwrap
        wrapped_lines = []
        for line in text.splitlines() or [""]:
            if len(line) <= max_chars:
                wrapped_lines.append(line)
            else:
                wrapped_lines.extend(textwrap.wrap(line, width=max_chars, break_long_words=True))
        return wrapped_lines

    def draw_paragraph(text: str):
        nonlocal y
        c.setFont("Helvetica", 10)
        for line in wrap_text(text):
            if y < 72:
                c.showPage()
                y = height - 72
            c.drawString(72, y, line)
            y -= 14
        y -= 8

    if draft.literature_summary:
        draw_paragraph("Literature Context (Auto-generated)")
        draw_paragraph(redact_phi(draft.literature_summary) if redact else draft.literature_summary)

    for q in questions:
        draw_paragraph(q.title)
        draw_paragraph(f"Prompt: {q.prompt}")
        ans = (draft.answers.get(q.category) or "").strip() or "[No response provided yet]"
        draw_paragraph(redact_phi(ans) if redact else ans)

    c.save()
    return output_path
